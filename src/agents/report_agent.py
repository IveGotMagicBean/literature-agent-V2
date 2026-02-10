"""
完全修复的报告生成Agent
1. Markdown报告打包成文件夹，使用相对路径
2. 正确的PDF生成
3. Word支持图片
"""

from pathlib import Path
from typing import Dict, Optional, List
import sys
import shutil

sys.path.insert(0, str(Path(__file__).parent))


class EnhancedReportAgent:
    """增强版报告生成Agent"""
    
    def __init__(self, llm, smart_agent):
        self.llm = llm
        self.smart_agent = smart_agent
    
    def generate(
        self, 
        report_type: str = "详细报告", 
        output_format: str = "PDF",
        include_figures: bool = True,
        max_figures: int = 5
    ) -> Optional[str]:
        """
        生成阅读报告
        
        Args:
            report_type: 报告类型（简要总结/详细报告/批判性分析）
            output_format: 输出格式（PDF/Word/Markdown）
            include_figures: 是否包含图片
            max_figures: 最多包含多少张图片
        
        Returns:
            报告文件路径
        """
        
        if not self.smart_agent.texts:
            raise ValueError("请先上传PDF文献")
        
        print(f"[Report Agent] 开始生成报告，格式: {output_format}")
        
        # 1. 生成报告内容
        print("[Report Agent] 生成报告内容...")
        report_content = self._generate_content(report_type)
        
        # 2. 提取关键图表
        figures = []
        if include_figures and self.smart_agent.figure_map:
            print("[Report Agent] 提取关键图表...")
            figures = self._extract_key_figures(max_figures)
        
        # 3. 格式化输出
        print(f"[Report Agent] 格式化输出为 {output_format}...")
        report_path = self._format_output(report_content, figures, output_format)
        
        print(f"[Report Agent] ✅ 报告生成完成: {report_path}")
        return report_path
    
    def _generate_content(self, report_type: str) -> Dict[str, str]:
        """生成报告内容（分段）"""
        
        if report_type == "简要总结":
            sections = self._generate_brief_summary()
        elif report_type == "详细报告":
            sections = self._generate_detailed_report()
        else:  # 批判性分析
            sections = self._generate_critical_analysis()
        
        return sections
    
    def _generate_brief_summary(self) -> Dict[str, str]:
        """生成简要总结"""
        
        text_sample = self.smart_agent.full_text[:5000]
        
        prompt = f"""请对以下论文进行简要总结（500字以内）：

文献内容：
{text_sample}

请用中文回答，包括：
1. 研究问题
2. 主要方法
3. 关键发现
4. 意义
"""
        
        summary = self.llm.chat(prompt)
        
        return {
            "摘要": summary
        }
    
    def _generate_detailed_report(self) -> Dict[str, str]:
        """生成详细报告"""
        
        text_sample = self.smart_agent.full_text[:8000]
        
        sections = {}
        
        # 基本信息
        prompt_basic = f"""请提取以下论文的基本信息：

文献内容：
{text_sample}

请用中文回答：
- 标题
- 作者及单位
- 发表信息（如果有）
"""
        sections["基本信息"] = self.llm.chat(prompt_basic)
        
        # 研究背景
        prompt_background = f"""请描述以下论文的研究背景与动机：

文献内容：
{text_sample}

请用中文回答：
- 研究领域
- 现有问题
- 研究意义
"""
        sections["研究背景与动机"] = self.llm.chat(prompt_background)
        
        # 研究方法
        prompt_method = f"""请描述以下论文的研究方法：

文献内容：
{text_sample}

请用中文回答：
- 技术路线
- 关键技术
- 创新点
"""
        sections["研究方法"] = self.llm.chat(prompt_method)
        
        # 实验与结果
        prompt_results = f"""请描述以下论文的实验与结果：

文献内容：
{text_sample}

请用中文回答：
- 实验设计
- 主要结果
- 数据分析
"""
        sections["实验与结果"] = self.llm.chat(prompt_results)
        
        # 讨论与结论
        prompt_conclusion = f"""请总结以下论文的讨论与结论：

文献内容：
{text_sample}

请用中文回答：
- 主要贡献
- 局限性
- 未来工作
"""
        sections["讨论与结论"] = self.llm.chat(prompt_conclusion)
        
        return sections
    
    def _generate_critical_analysis(self) -> Dict[str, str]:
        """生成批判性分析"""
        
        text_sample = self.smart_agent.full_text[:8000]
        
        sections = {}
        
        prompt_template = """请对以下论文进行批判性分析：

文献内容：
{text}

请用中文分析：
{aspect}
"""
        
        aspects = {
            "研究问题的合理性": "问题是否重要？问题定义是否清晰？",
            "方法的有效性": "方法是否合适？是否有更好的替代方法？",
            "实验的严谨性": "实验设计是否合理？数据是否充分？对比是否公平？",
            "结论的可靠性": "结论是否有数据支撑？是否存在过度推广？",
            "创新性评估": "真正的创新点在哪？是否只是工程改进？"
        }
        
        for section, aspect in aspects.items():
            sections[section] = self.llm.chat(
                prompt_template.format(text=text_sample, aspect=aspect)
            )
        
        return sections
    
    def _extract_key_figures(self, max_figures: int = 5) -> List[Dict]:
        """提取关键图表"""
        
        key_figures = []
        
        for fig_num, fig_data in list(self.smart_agent.figure_map.items())[:max_figures]:
            if not fig_data or "path" not in fig_data:
                continue
            
            # 确保图片路径存在
            fig_path = Path(fig_data["path"])
            if not fig_path.exists():
                print(f"[Report Agent] 警告: 图片不存在 {fig_path}")
                continue
            
            # 获取图表描述
            description = self._get_figure_summary(fig_num)
            
            key_figures.append({
                "number": fig_num,
                "path": str(fig_path.absolute()),
                "description": description
            })
        
        return key_figures
    
    def _get_figure_summary(self, fig_num: int) -> str:
        """获取图表简介"""
        
        descriptions = []
        for text_data in self.smart_agent.texts:
            refs = self.smart_agent.ref_extractor.extract_references(text_data["content"])
            for ref in refs:
                if ref["figure"] == fig_num:
                    context = self.smart_agent.ref_extractor.get_context(
                        text_data["content"], ref, 150
                    )
                    descriptions.append(context)
                    if len(descriptions) >= 1:
                        break
            if descriptions:
                break
        
        return descriptions[0] if descriptions else f"Figure {fig_num}"
    
    def _format_output(
        self, 
        content: Dict[str, str], 
        figures: List[Dict],
        output_format: str
    ) -> str:
        """格式化输出"""
        
        from datetime import datetime
        
        output_dir = Path(self.smart_agent.data_dir)
        output_dir.mkdir(exist_ok=True)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # 根据格式调用不同的生成方法
        if output_format.upper() == "PDF":
            return self._create_pdf(content, figures, output_dir, timestamp)
        elif output_format.upper() == "WORD" or output_format.upper() == "DOCX":
            return self._create_word(content, figures, output_dir, timestamp)
        elif output_format.upper() == "MARKDOWN" or output_format.upper() == "MD":
            return self._create_markdown_package(content, figures, output_dir, timestamp)
        else:
            # 默认PDF
            print(f"[Report Agent] 未知格式 '{output_format}'，使用PDF")
            return self._create_pdf(content, figures, output_dir, timestamp)
    
    def _create_pdf(
        self, 
        content: Dict[str, str], 
        figures: List[Dict],
        output_dir: Path,
        timestamp: str
    ) -> str:
        """创建PDF报告（带图片）"""
        
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            
            print("[Report Agent] 使用reportlab生成PDF...")
            
            output_path = output_dir / f"report_{timestamp}.pdf"
            
            # 创建PDF文档
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18,
            )
            
            # 获取样式
            styles = getSampleStyleSheet()
            
            # 创建自定义样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor='darkblue',
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor='darkblue',
                spaceAfter=12,
                spaceBefore=12
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=11,
                alignment=TA_JUSTIFY,
                spaceAfter=12
            )
            
            # 构建内容
            story = []
            
            # 标题
            story.append(Paragraph("学术论文阅读报告", title_style))
            story.append(Spacer(1, 12))
            
            # 时间戳
            story.append(Paragraph(f"生成时间: {timestamp}", styles['Normal']))
            story.append(Spacer(1, 24))
            
            # 添加各章节内容
            for section_title, section_text in content.items():
                # 章节标题
                story.append(Paragraph(section_title, heading_style))
                
                # 章节内容（处理换行）
                paragraphs = section_text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        # 清理文本中可能的特殊字符
                        clean_para = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(clean_para, body_style))
                
                story.append(Spacer(1, 12))
            
            # 添加图片章节
            if figures:
                story.append(PageBreak())
                story.append(Paragraph("关键图表", heading_style))
                story.append(Spacer(1, 12))
                
                for fig in figures:
                    # 图片标题
                    fig_title = f"Figure {fig['number']}"
                    story.append(Paragraph(fig_title, heading_style))
                    story.append(Spacer(1, 6))
                    
                    # 添加图片
                    try:
                        img_path = Path(fig['path'])
                        if img_path.exists():
                            # 创建图片对象（限制宽度）
                            img = Image(str(img_path), width=5*inch, height=None)
                            img.hAlign = 'CENTER'
                            story.append(img)
                            story.append(Spacer(1, 6))
                            
                            print(f"[Report Agent] ✓ 图片已添加到PDF: {img_path.name}")
                        else:
                            story.append(Paragraph(f"[图片文件不存在: {img_path}]", styles['Normal']))
                    except Exception as e:
                        print(f"[Report Agent] 添加图片到PDF失败: {e}")
                        story.append(Paragraph(f"[无法添加图片: {e}]", styles['Normal']))
                    
                    # 图片描述
                    if fig.get('description'):
                        desc_style = ParagraphStyle(
                            'FigureDesc',
                            parent=styles['Normal'],
                            fontSize=10,
                            alignment=TA_CENTER,
                            textColor='gray',
                            fontName='Times-Italic'
                        )
                        clean_desc = fig['description'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(clean_desc, desc_style))
                    
                    story.append(Spacer(1, 24))
            
            # 生成PDF
            doc.build(story)
            
            print(f"[Report Agent] ✅ PDF生成成功: {output_path}")
            return str(output_path)
            
        except ImportError as e:
            print(f"[Report Agent] reportlab未安装: {e}")
            print("[Report Agent] 降级到Markdown...")
            return self._create_markdown_package(content, figures, output_dir, timestamp)
        except Exception as e:
            print(f"[Report Agent] PDF生成失败: {e}")
            import traceback
            traceback.print_exc()
            print("[Report Agent] 降级到Markdown...")
            return self._create_markdown_package(content, figures, output_dir, timestamp)
    
    def _create_word(
        self, 
        content: Dict[str, str], 
        figures: List[Dict],
        output_dir: Path,
        timestamp: str
    ) -> str:
        """创建Word报告（带图片）"""
        
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            print("[Report Agent] 使用python-docx生成Word...")
            
            output_path = output_dir / f"report_{timestamp}.docx"
            
            doc = Document()
            
            # 添加标题
            title = doc.add_heading("学术论文阅读报告", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加生成时间
            doc.add_paragraph(f"生成时间: {timestamp}")
            doc.add_paragraph()
            
            # 添加各部分内容
            for section, text in content.items():
                doc.add_heading(section, level=1)
                
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        if para.strip().startswith('- '):
                            doc.add_paragraph(para.strip()[2:], style='List Bullet')
                        else:
                            doc.add_paragraph(para.strip())
                
                doc.add_paragraph()
            
            # 添加图片章节
            if figures:
                doc.add_heading("关键图表", level=1)
                
                for fig in figures:
                    doc.add_heading(f"Figure {fig['number']}", level=2)
                    
                    try:
                        img_path = Path(fig['path'])
                        if img_path.exists():
                            doc.add_picture(str(img_path), width=Inches(5.5))
                            
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            print(f"[Report Agent] ✓ 图片已添加到Word: {img_path.name}")
                        else:
                            doc.add_paragraph(f"[图片文件不存在: {img_path}]")
                    except Exception as e:
                        print(f"[Report Agent] 添加图片到Word失败: {e}")
                        doc.add_paragraph(f"[无法添加图片: {e}]")
                    
                    if fig.get('description'):
                        p = doc.add_paragraph(fig['description'])
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.size = Pt(10)
                            run.font.italic = True
                    
                    doc.add_paragraph()
            
            doc.save(str(output_path))
            
            print(f"[Report Agent] ✅ Word生成成功: {output_path}")
            return str(output_path)
            
        except ImportError as e:
            print(f"[Report Agent] python-docx未安装: {e}")
            print("[Report Agent] 降级到Markdown...")
            return self._create_markdown_package(content, figures, output_dir, timestamp)
        except Exception as e:
            print(f"[Report Agent] Word生成失败: {e}")
            import traceback
            traceback.print_exc()
            print("[Report Agent] 降级到Markdown...")
            return self._create_markdown_package(content, figures, output_dir, timestamp)
    
    def _create_markdown_package(
        self, 
        content: Dict[str, str], 
        figures: List[Dict],
        output_dir: Path,
        timestamp: str
    ) -> str:
        """创建Markdown报告包（文件夹，包含图片，使用相对路径）"""
        
        print("[Report Agent] 生成Markdown报告包...")
        
        # 创建报告文件夹
        report_folder = output_dir / f"report_{timestamp}"
        report_folder.mkdir(exist_ok=True)
        
        # 创建images子文件夹
        images_folder = report_folder / "images"
        images_folder.mkdir(exist_ok=True)
        
        # 复制图片到报告文件夹
        copied_figures = []
        for fig in figures:
            try:
                src_path = Path(fig['path'])
                if src_path.exists():
                    # 使用figure编号命名
                    dst_filename = f"figure_{fig['number']}{src_path.suffix}"
                    dst_path = images_folder / dst_filename
                    
                    shutil.copy2(src_path, dst_path)
                    
                    copied_figures.append({
                        "number": fig["number"],
                        "path": f"images/{dst_filename}",  # 相对路径
                        "description": fig["description"]
                    })
                    
                    print(f"[Report Agent] ✓ 图片已复制: {dst_filename}")
            except Exception as e:
                print(f"[Report Agent] 复制图片失败: {e}")
        
        # 生成Markdown内容（使用相对路径）
        md_content = "# 学术论文阅读报告\n\n"
        md_content += f"> 生成时间: {timestamp}\n\n"
        md_content += "---\n\n"
        
        # 添加各部分内容
        for section, text in content.items():
            md_content += f"## {section}\n\n{text}\n\n---\n\n"
        
        # 添加图片（使用相对路径）
        if copied_figures:
            md_content += "## 关键图表\n\n"
            for fig in copied_figures:
                md_content += f"### Figure {fig['number']}\n\n"
                md_content += f"![Figure {fig['number']}]({fig['path']})\n\n"
                md_content += f"{fig['description']}\n\n---\n\n"
        
        # 保存Markdown文件
        md_path = report_folder / "report.md"
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        print(f"[Report Agent] ✅ Markdown报告包生成成功: {report_folder}")
        print(f"[Report Agent]    - Markdown文件: report.md")
        print(f"[Report Agent]    - 图片文件夹: images/ ({len(copied_figures)}张图片)")
        
        # 返回文件夹路径
        return str(report_folder)


__all__ = ["EnhancedReportAgent"]
